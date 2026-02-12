"""
Author Studio - Manuscript Converter
Converts Markdown to Word with proper formatting
"""

import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from datetime import datetime


class MarkdownToWord:
    """Converts Markdown files to properly formatted Word documents."""
    
    @staticmethod
    def parse_markdown(md_text: str) -> Dict:
        """Parse markdown into structured content."""
        lines = md_text.split('\n')
        
        content = {
            'title': None,
            'subtitle': None,
            'author': None,
            'sections': [],
            'metadata': {}
        }
        
        current_section = None
        current_content = []
        
        # Check for YAML frontmatter
        if lines and lines[0].strip() == '---':
            frontmatter_end = -1
            for i, line in enumerate(lines[1:], 1):
                if line.strip() == '---':
                    frontmatter_end = i
                    break
            
            if frontmatter_end > 0:
                frontmatter = '\n'.join(lines[1:frontmatter_end])
                content['metadata'] = MarkdownToWord._parse_frontmatter(frontmatter)
                lines = lines[frontmatter_end + 1:]
                
                # Extract title/author from frontmatter
                content['title'] = content['metadata'].get('title')
                content['author'] = content['metadata'].get('author')
                content['subtitle'] = content['metadata'].get('subtitle')
        
        for line in lines:
            stripped = line.strip()
            
            # Headers
            if stripped.startswith('# ') and not content['title']:
                content['title'] = stripped[2:].strip()
            elif stripped.startswith('# '):
                # Save previous section
                if current_section:
                    current_section['content'] = '\n'.join(current_content)
                    content['sections'].append(current_section)
                
                current_section = {
                    'level': 1,
                    'title': stripped[2:].strip(),
                    'content': ''
                }
                current_content = []
            elif stripped.startswith('## '):
                if current_section:
                    current_section['content'] = '\n'.join(current_content)
                    content['sections'].append(current_section)
                
                current_section = {
                    'level': 2,
                    'title': stripped[3:].strip(),
                    'content': ''
                }
                current_content = []
            elif stripped.startswith('### '):
                if current_section:
                    current_section['content'] = '\n'.join(current_content)
                    content['sections'].append(current_section)
                
                current_section = {
                    'level': 3,
                    'title': stripped[4:].strip(),
                    'content': ''
                }
                current_content = []
            else:
                current_content.append(line)
        
        # Save last section
        if current_section:
            current_section['content'] = '\n'.join(current_content)
            content['sections'].append(current_section)
        elif current_content:
            content['sections'].append({
                'level': 0,
                'title': None,
                'content': '\n'.join(current_content)
            })
        
        return content
    
    @staticmethod
    def _parse_frontmatter(text: str) -> Dict:
        """Parse YAML-like frontmatter."""
        metadata = {}
        for line in text.split('\n'):
            if ':' in line:
                key, value = line.split(':', 1)
                metadata[key.strip().lower()] = value.strip().strip('"\'')
        return metadata
    
    @staticmethod
    def convert_to_docx(md_text: str, output_path: str, 
                       title: str = None, author: str = None) -> Dict:
        """
        Convert Markdown text to Word document.
        
        Args:
            md_text: Markdown content
            output_path: Path to save the .docx file
            title: Override title
            author: Override author
            
        Returns:
            Dict with success status and info
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
        except ImportError:
            return {'success': False, 'error': 'python-docx not installed'}
        
        # Parse markdown
        content = MarkdownToWord.parse_markdown(md_text)
        
        # Override with provided values
        if title:
            content['title'] = title
        if author:
            content['author'] = author
        
        # Create document
        doc = Document()
        
        # Set up styles
        styles = doc.styles
        
        # Title page
        if content['title']:
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(content['title'])
            title_run.bold = True
            title_run.font.size = Pt(28)
            
            if content.get('subtitle'):
                doc.add_paragraph()
                subtitle_para = doc.add_paragraph()
                subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                subtitle_run = subtitle_para.add_run(content['subtitle'])
                subtitle_run.font.size = Pt(18)
            
            if content.get('author'):
                doc.add_paragraph()
                doc.add_paragraph()
                author_para = doc.add_paragraph()
                author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                author_run = author_para.add_run(f"by {content['author']}")
                author_run.font.size = Pt(14)
            
            # Page break after title
            doc.add_page_break()
        
        # Add sections
        for section in content['sections']:
            if section['title']:
                level = section['level']
                
                if level == 1:
                    heading = doc.add_heading(section['title'], level=1)
                elif level == 2:
                    heading = doc.add_heading(section['title'], level=2)
                elif level == 3:
                    heading = doc.add_heading(section['title'], level=3)
            
            # Add content
            if section['content']:
                paragraphs = section['content'].split('\n\n')
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if not para_text:
                        continue
                    
                    # Check for bullet points
                    if para_text.startswith('- ') or para_text.startswith('* '):
                        items = re.split(r'\n[-*]\s+', para_text)
                        for item in items:
                            item = item.lstrip('- *').strip()
                            if item:
                                p = doc.add_paragraph(item, style='List Bullet')
                    
                    # Check for numbered list
                    elif re.match(r'^\d+\.\s', para_text):
                        items = re.split(r'\n\d+\.\s+', para_text)
                        for item in items:
                            item = re.sub(r'^\d+\.\s*', '', item).strip()
                            if item:
                                p = doc.add_paragraph(item, style='List Number')
                    
                    # Check for blockquote
                    elif para_text.startswith('>'):
                        quote_text = para_text.lstrip('> ').strip()
                        p = doc.add_paragraph(quote_text)
                        p.paragraph_format.left_indent = Inches(0.5)
                        p.runs[0].italic = True
                    
                    # Regular paragraph
                    else:
                        # Process inline formatting
                        p = doc.add_paragraph()
                        MarkdownToWord._add_formatted_text(p, para_text)
        
        # Save document
        doc.save(output_path)
        
        return {
            'success': True,
            'output_path': output_path,
            'title': content['title'],
            'sections': len(content['sections']),
            'word_count': len(md_text.split())
        }
    
    @staticmethod
    def _add_formatted_text(paragraph, text: str):
        """Add text with inline markdown formatting to a paragraph."""
        # Simple pattern matching for bold, italic, code
        pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|__[^_]+__|_[^_]+_)'
        
        parts = re.split(pattern, text)
        
        for part in parts:
            if not part:
                continue
            
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('__') and part.endswith('__'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('_') and part.endswith('_'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Courier New'
            else:
                paragraph.add_run(part)


class WordToMarkdown:
    """Converts Word documents to Markdown."""
    
    @staticmethod
    def convert(docx_path: str) -> Dict:
        """
        Convert Word document to Markdown.
        
        Args:
            docx_path: Path to .docx file
            
        Returns:
            Dict with success status and markdown content
        """
        try:
            from docx import Document
        except ImportError:
            return {'success': False, 'error': 'python-docx not installed'}
        
        try:
            doc = Document(docx_path)
        except Exception as e:
            return {'success': False, 'error': str(e)}
        
        md_lines = []
        
        for para in doc.paragraphs:
            style_name = para.style.name.lower() if para.style else ''
            text = para.text.strip()
            
            if not text:
                md_lines.append('')
                continue
            
            # Convert headings
            if 'heading 1' in style_name or 'title' in style_name:
                md_lines.append(f'# {text}')
            elif 'heading 2' in style_name:
                md_lines.append(f'## {text}')
            elif 'heading 3' in style_name:
                md_lines.append(f'### {text}')
            elif 'heading 4' in style_name:
                md_lines.append(f'#### {text}')
            elif 'list bullet' in style_name:
                md_lines.append(f'- {text}')
            elif 'list number' in style_name:
                md_lines.append(f'1. {text}')
            else:
                # Check for inline formatting
                formatted_text = WordToMarkdown._extract_formatting(para)
                md_lines.append(formatted_text)
        
        markdown = '\n\n'.join(md_lines)
        
        # Clean up extra blank lines
        markdown = re.sub(r'\n{3,}', '\n\n', markdown)
        
        return {
            'success': True,
            'markdown': markdown,
            'word_count': len(markdown.split())
        }
    
    @staticmethod
    def _extract_formatting(paragraph) -> str:
        """Extract formatting from paragraph runs."""
        result = []
        
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
            
            if run.bold and run.italic:
                result.append(f'***{text}***')
            elif run.bold:
                result.append(f'**{text}**')
            elif run.italic:
                result.append(f'*{text}*')
            else:
                result.append(text)
        
        return ''.join(result)


class ManuscriptFormatter:
    """Formats manuscripts for publishing."""
    
    # Standard book formatting
    BOOK_SETTINGS = {
        'trim_size': '6x9',
        'margins': {
            'top': 0.75,
            'bottom': 0.75,
            'inside': 0.875,  # Gutter
            'outside': 0.625
        },
        'font': 'Garamond',
        'font_size': 11,
        'line_spacing': 1.15,
        'paragraph_indent': 0.3,
        'chapter_start_new_page': True
    }
    
    @staticmethod
    def format_for_print(docx_path: str, output_path: str = None,
                        trim_size: str = '6x9') -> Dict:
        """
        Format a Word document for print publishing.
        
        Args:
            docx_path: Path to input .docx
            output_path: Path to save formatted .docx (default: adds _formatted)
            trim_size: Book size ('6x9', '5.5x8.5', '5x8')
            
        Returns:
            Dict with success status
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return {'success': False, 'error': 'python-docx not installed'}
        
        if not output_path:
            path = Path(docx_path)
            output_path = str(path.parent / f"{path.stem}_formatted{path.suffix}")
        
        try:
            doc = Document(docx_path)
        except Exception as e:
            return {'success': False, 'error': str(e)}
        
        # Set page size based on trim
        sizes = {
            '6x9': (6, 9),
            '5.5x8.5': (5.5, 8.5),
            '5x8': (5, 8),
            '8.5x11': (8.5, 11)
        }
        
        width, height = sizes.get(trim_size, (6, 9))
        
        for section in doc.sections:
            section.page_width = Inches(width)
            section.page_height = Inches(height)
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.875)  # Inside/gutter
            section.right_margin = Inches(0.625)  # Outside
        
        # Format paragraphs
        for para in doc.paragraphs:
            style_name = para.style.name.lower() if para.style else ''
            
            # Skip headings
            if 'heading' in style_name or 'title' in style_name:
                continue
            
            # Set paragraph formatting
            para.paragraph_format.first_line_indent = Inches(0.3)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.15
            
            # Set font
            for run in para.runs:
                run.font.name = 'Garamond'
                run.font.size = Pt(11)
        
        doc.save(output_path)
        
        return {
            'success': True,
            'output_path': output_path,
            'trim_size': trim_size,
            'page_size': f"{width}x{height} inches"
        }
    
    @staticmethod
    def add_front_matter(docx_path: str, 
                        title: str,
                        author: str,
                        copyright_year: int = None,
                        dedication: str = None,
                        isbn: str = None) -> Dict:
        """Add front matter pages to a manuscript."""
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return {'success': False, 'error': 'python-docx not installed'}
        
        if not copyright_year:
            copyright_year = datetime.now().year
        
        try:
            doc = Document(docx_path)
        except Exception as e:
            return {'success': False, 'error': str(e)}
        
        # We need to insert at the beginning, which is tricky in python-docx
        # For now, return instructions
        
        front_matter = f"""
FRONT MATTER TO ADD:

=== TITLE PAGE ===
{title}
by {author}

=== COPYRIGHT PAGE ===
Copyright © {copyright_year} {author}
All rights reserved.
{f'ISBN: {isbn}' if isbn else ''}

No part of this publication may be reproduced, distributed, or transmitted 
in any form or by any means without the prior written permission of the publisher.

{f'=== DEDICATION ===' if dedication else ''}
{dedication if dedication else ''}
"""
        
        return {
            'success': True,
            'front_matter': front_matter,
            'note': 'Add this front matter manually at the beginning of your document'
        }


# Convenience functions
def convert_md_to_docx(md_path: str, output_path: str = None) -> Dict:
    """Convert a Markdown file to Word document."""
    md_path = Path(md_path)
    
    if not md_path.exists():
        return {'success': False, 'error': 'File not found'}
    
    if not output_path:
        output_path = str(md_path.parent / f"{md_path.stem}.docx")
    
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()
    
    return MarkdownToWord.convert_to_docx(md_text, output_path)


def convert_docx_to_md(docx_path: str, output_path: str = None) -> Dict:
    """Convert a Word document to Markdown."""
    result = WordToMarkdown.convert(docx_path)
    
    if result['success'] and output_path:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(result['markdown'])
        result['output_path'] = output_path
    
    return result
